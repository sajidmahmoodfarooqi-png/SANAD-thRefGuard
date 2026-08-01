"""Apply a Style Profile's formatting to a researcher's thesis .docx.

The user's actual document goes in; the same document comes back reformatted to
the university's rules — body font/size/line-spacing, page margins, heading fonts,
and a hanging indent on the reference list — with **not one word of prose
changed**. That last part is a hard guarantee, enforced two ways:

  * this module only ever writes *style and format* properties (Style.font,
    ParagraphFormat, Section margins). It never reads or assigns paragraph/run
    *text*. There is no code path here that could alter wording.
  * `text_fingerprint()` captures every paragraph + table cell string; the caller
    asserts it is byte-identical before and after, so a regression can't slip
    through silently.

The whole routine runs inside `offline.no_network()`, so the document can never
leave the machine while being processed (see offline.py for why that matters).
"""
from __future__ import annotations

import io

from . import offline


def text_fingerprint(data: bytes) -> list[str]:
    """Every run of user text in the document, in order — the thing that must
    NOT change. Used to prove prose was untouched."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    out: list[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                out.extend(p.text for p in cell.paragraphs)
    return out


def _ensure_numbering_part(doc):
    """Return the document's <w:numbering> element, creating the numbering part
    if it has none. A document that never contained a list has no numbering.xml
    at all, and python-docx will not create one for you (NumberingPart.new()
    raises NotImplementedError) — so a clean thesis got NO heading numbers. We
    build a minimal numbering part and attach it. Still only a *definition*; it
    never writes prose."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc_part = doc.part
    try:
        return doc_part.numbering_part.element
    except (NotImplementedError, KeyError):
        pass
    from docx.opc.packuri import PackURI
    from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
    from docx.parts.numbering import NumberingPart

    element = parse_xml(f'<w:numbering {nsdecls("w")}/>')
    part = NumberingPart(PackURI("/word/numbering.xml"), CT.WML_NUMBERING,
                         element, doc_part.package)
    doc_part.relate_to(part, RT.NUMBERING)
    return element


def _add_multilevel_numbering(doc):
    """Inject Word's canonical heading multilevel list (1 / 1.1 / 1.1.1) into the
    numbering part and return its numId. Each level names its heading style via
    <w:pStyle> — exactly how Word's own 'Multilevel List → Heading' scheme is
    built — which is what makes the numbers actually render (and render on the
    heading's own line). Raw OOXML, but only a numbering *definition*: no text."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn, nsdecls
    import secrets

    numbering = _ensure_numbering_part(doc)
    aids = [int(e.get(qn("w:abstractNumId"))) for e in numbering.findall(qn("w:abstractNum"))]
    nids = [int(e.get(qn("w:numId"))) for e in numbering.findall(qn("w:num"))]
    aid = (max(aids) + 1) if aids else 0
    nid = (max(nids) + 1) if nids else 1
    lvls = ""
    # (ilvl, heading style id, level text). Heading 1/2/3 -> 1 / 1.1 / 1.1.1
    for ilvl, pstyle, fmt in ((0, "Heading1", "%1"), (1, "Heading2", "%1.%2"),
                              (2, "Heading3", "%1.%2.%3")):
        # element order matters (CT_Lvl schema sequence): start, numFmt, pStyle,
        # suff, lvlText, lvlJc, pPr. suff=space avoids the tab-stop that made a
        # heading appear to "drift" to the end of the previous line.
        lvls += (f'<w:lvl w:ilvl="{ilvl}">'
                 f'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
                 f'<w:pStyle w:val="{pstyle}"/><w:suff w:val="space"/>'
                 f'<w:lvlText w:val="{fmt}"/><w:lvlJc w:val="left"/>'
                 f'<w:pPr><w:ind w:left="0" w:firstLine="0"/></w:pPr></w:lvl>')
    abs_el = parse_xml(f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{aid}">'
                       f'<w:nsid w:val="{secrets.token_hex(4).upper()}"/>'
                       f'<w:multiLevelType w:val="multilevel"/>{lvls}</w:abstractNum>')
    num_el = parse_xml(f'<w:num {nsdecls("w")} w:numId="{nid}"><w:abstractNumId w:val="{aid}"/></w:num>')
    first_num = numbering.find(qn("w:num"))
    first_num.addprevious(abs_el) if first_num is not None else numbering.append(abs_el)
    numbering.append(num_el)
    return nid


def _link_heading_numbering(doc, style_name, ilvl, numid):
    """Bind a heading style to a level of the numbering list, so every paragraph
    using that style is auto-numbered by Word (the number is rendered, not typed)."""
    from docx.oxml.ns import qn
    pPr = doc.styles[style_name].element.get_or_add_pPr()
    for old in pPr.findall(qn("w:numPr")):
        pPr.remove(old)
    numPr = pPr.get_or_add_numPr()
    numPr.get_or_add_ilvl().val = ilvl
    numPr.get_or_add_numId().val = numid


def _apply_headings(doc, headings, body_font, applied):
    """Set Title / Heading 1-3 fonts + sizes and, if requested, the 1/1.1/1.1.1
    multilevel numbering. Falls back to the body font for headings with none set."""
    from docx.shared import Pt

    levels = {"Title": headings.get("title") or {}, "Heading 1": headings.get("h1") or {},
              "Heading 2": headings.get("h2") or {}, "Heading 3": headings.get("h3") or {}}
    for name, spec in levels.items():
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        font = spec.get("font") or body_font
        if font:
            style.font.name = font
        if spec.get("size_pt") is not None:
            style.font.size = Pt(float(spec["size_pt"]))
            applied.append(f"{name} size → {spec['size_pt']} pt")

    if headings.get("numbered"):
        try:
            nid = _add_multilevel_numbering(doc)
            for name, ilvl in (("Heading 1", 0), ("Heading 2", 1), ("Heading 3", 2)):
                if name in [s.name for s in doc.styles]:
                    _link_heading_numbering(doc, name, ilvl, nid)
            applied.append("Heading numbering → 1 / 1.1 / 1.1.1 (Heading 1/2/3)")
        except Exception as exc:  # numbering is best-effort; never fail the whole reformat
            applied.append(f"Heading numbering: skipped ({exc.__class__.__name__})")


def _apply_caption(doc, caption, applied):
    """Style Word's built-in Caption style (used for figure/table captions):
    font, size, and italic/regular. We only style *how* a caption looks — we
    never insert captions or move them, so no prose is touched. Word decides
    placement (below a figure, above a table) when the caption is inserted."""
    from docx.shared import Pt

    if not caption or not (caption.get("font") or caption.get("size_pt") is not None
                           or caption.get("italic") is not None):
        return
    try:
        style = doc.styles["Caption"]
    except KeyError:
        applied.append("Caption style: skipped (document has no 'Caption' style yet — "
                       "insert one caption in Word first, then reformat)")
        return
    if caption.get("font"):
        style.font.name = caption["font"]
    if caption.get("size_pt") is not None:
        style.font.size = Pt(float(caption["size_pt"]))
    if caption.get("italic") is not None:
        style.font.italic = bool(caption["italic"])
    bits = []
    if caption.get("font"):
        bits.append(caption["font"])
    if caption.get("size_pt") is not None:
        bits.append(f"{caption['size_pt']} pt")
    if caption.get("italic"):
        bits.append("italic")
    applied.append("Caption style → " + (", ".join(bits) or "updated"))


def apply_profile_to_docx(data: bytes, profile: dict) -> dict:
    """Reformat the .docx to the profile. Returns {data (bytes), applied (list)}.
    Raises ValueError on a non-.docx or corrupt file. Never touches prose."""
    from docx import Document
    from docx.shared import Cm, Pt

    ps = profile.get("paragraph_style") or {}
    ds = profile.get("document_structure") or {}
    font = ps.get("font_family")
    size = ps.get("font_size_pt")
    spacing = ps.get("line_spacing")
    margin_cm = ds.get("margin_cm") if ds.get("enabled") else None
    binding_cm = ds.get("binding_margin_cm") if ds.get("enabled") else None
    binding_side = (ds.get("binding_side") or "left").lower()
    indent_cm = ps.get("bibliography_hanging_indent_cm")
    caption = ds.get("caption") or {}
    applied: list[str] = []

    with offline.no_network():          # the document cannot leave the machine
        try:
            before = text_fingerprint(data)
            doc = Document(io.BytesIO(data))
        except ValueError:
            raise
        except Exception as exc:
            raise ValueError("that doesn't look like a valid .docx file") from exc

        # body defaults live on the Normal style
        normal = doc.styles["Normal"]
        if font:
            normal.font.name = font
            applied.append(f"Body font → {font}")
        if size:
            normal.font.size = Pt(float(size))
            applied.append(f"Body size → {size} pt")
        if spacing:
            normal.paragraph_format.line_spacing = float(spacing)
            applied.append(f"Line spacing → {spacing}")

        # page margins on every section. A larger binding-side (gutter) margin is
        # applied on top for thesis binding — e.g. 1 inch all round, 1.5 inch on
        # the left (binding) edge.
        if margin_cm:
            for section in doc.sections:
                section.top_margin = section.bottom_margin = Cm(float(margin_cm))
                section.left_margin = section.right_margin = Cm(float(margin_cm))
            applied.append(f"Margins → {margin_cm} cm")
        if binding_cm:
            attr = {"left": "left_margin", "right": "right_margin",
                    "top": "top_margin", "bottom": "bottom_margin"}.get(binding_side, "left_margin")
            for section in doc.sections:
                setattr(section, attr, Cm(float(binding_cm)))
            applied.append(f"Binding margin ({binding_side}) → {binding_cm} cm")

        # Word Styles gallery: Title / Heading 1-3 fonts + sizes, and the
        # standard 1 / 1.1 / 1.1.1 multilevel numbering when requested
        _apply_headings(doc, ds.get("headings") or {}, font, applied)

        # caption style (for figures/tables): style the built-in Caption style.
        # Placement (figures below, tables above) is chosen when you insert the
        # caption in Word — this only sets how captions *look*.
        _apply_caption(doc, caption, applied)

        # reference list hanging indent: apply to the built-in Bibliography style
        # if the document uses it (positive left indent + equal negative first line)
        if indent_cm is not None:
            try:
                bib = doc.styles["Bibliography"]
                bib.paragraph_format.left_indent = Cm(float(indent_cm))
                bib.paragraph_format.first_line_indent = Cm(-float(indent_cm))
                applied.append(f"Reference hanging indent → {indent_cm} cm")
            except KeyError:
                applied.append("Reference hanging indent: skipped (document has no "
                               "'Bibliography' style — apply that style to your reference list)")

        out = io.BytesIO()
        doc.save(out)
        after = text_fingerprint(out.getvalue())

    if before != after:  # the guarantee: prose is byte-identical
        raise RuntimeError("aborted: formatting would have altered document text "
                           "(this should be impossible — no text is written)")
    return {"data": out.getvalue(), "applied": applied}
