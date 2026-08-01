"""Thesis .docx reformatting: applies the profile, never changes prose, stays offline."""
import base64
import io

import pytest
from fastapi.testclient import TestClient

from sanad_core import docformat
from sanad_core.server import create_app

PROSE = [
    "Chapter 1: Introduction",
    "This is the researcher's own sentence and it must never change.",
    "A second paragraph with specific wording: azimuth, provenance, 1948.",
]


def _thesis_docx():
    from docx import Document
    doc = Document()
    for line in PROSE:
        doc.add_paragraph(line)
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


PROFILE = {
    "paragraph_style": {"font_family": "Georgia", "font_size_pt": 13, "line_spacing": 2.0,
                        "bibliography_hanging_indent_cm": 1.27},
    "document_structure": {"enabled": True, "margin_cm": 2.54},
}


def test_reformat_applies_and_never_touches_prose():
    data = _thesis_docx()
    before = docformat.text_fingerprint(data)
    result = docformat.apply_profile_to_docx(data, PROFILE)
    after = docformat.text_fingerprint(result["data"])
    assert before == after == PROSE          # every word identical, before and after
    # formatting really landed on the Normal style
    from docx import Document
    from docx.shared import Pt
    doc = Document(io.BytesIO(result["data"]))
    normal = doc.styles["Normal"]
    assert normal.font.name == "Georgia"
    assert normal.font.size == Pt(13)
    assert abs(doc.sections[0].top_margin.cm - 2.54) < 0.01
    assert any("Body font" in a for a in result["applied"])


def test_reformat_rejects_non_docx():
    with pytest.raises(ValueError):
        docformat.apply_profile_to_docx(b"not a docx at all", PROFILE)


HEADING_PROSE = [
    "Introduction",          # Heading 1
    "Background",            # Heading 2
    "Prior work",           # Heading 3
    "The researcher's own body sentence, untouched.",
]


def _thesis_with_headings():
    from docx import Document
    doc = Document()
    doc.add_paragraph(HEADING_PROSE[0], style="Heading 1")
    doc.add_paragraph(HEADING_PROSE[1], style="Heading 2")
    doc.add_paragraph(HEADING_PROSE[2], style="Heading 3")
    doc.add_paragraph(HEADING_PROSE[3])
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


HEADING_PROFILE = {
    "paragraph_style": {"font_family": "Georgia", "font_size_pt": 12},
    "document_structure": {"enabled": True, "headings": {
        "numbered": True,
        "title": {"font": "Cambria", "size_pt": 26},
        "h1": {"font": "Cambria", "size_pt": 16},
        "h2": {"font": "Cambria", "size_pt": 14},
        "h3": {"size_pt": 12},           # no font → inherits body font
    }}}


def test_heading_fonts_and_sizes_applied_prose_untouched():
    from docx import Document
    from docx.shared import Pt
    data = _thesis_with_headings()
    before = docformat.text_fingerprint(data)
    result = docformat.apply_profile_to_docx(data, HEADING_PROFILE)
    after = docformat.text_fingerprint(result["data"])
    assert before == after == HEADING_PROSE      # heading + body text all identical

    doc = Document(io.BytesIO(result["data"]))
    assert doc.styles["Title"].font.name == "Cambria"
    assert doc.styles["Title"].font.size == Pt(26)
    assert doc.styles["Heading 1"].font.size == Pt(16)
    assert doc.styles["Heading 2"].font.size == Pt(14)
    # h3 had no font of its own → falls back to the body font
    assert doc.styles["Heading 3"].font.name == "Georgia"
    assert doc.styles["Heading 3"].font.size == Pt(12)


def test_multilevel_numbering_injected_and_linked():
    from docx import Document
    from docx.oxml.ns import qn
    data = _thesis_with_headings()
    result = docformat.apply_profile_to_docx(data, HEADING_PROFILE)
    doc = Document(io.BytesIO(result["data"]))

    # the 1 / 1.1 / 1.1.1 level texts exist in the numbering part
    numbering = doc.part.numbering_part.element
    texts = [e.get(qn("w:val")) for e in numbering.iter(qn("w:lvlText"))]
    assert "%1" in texts and "%1.%2" in texts and "%1.%2.%3" in texts

    # each heading style points at the right level of that same numId
    def linked(style_name):
        pPr = doc.styles[style_name].element.find(qn("w:pPr"))
        numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
        assert numPr is not None, f"{style_name} not linked to numbering"
        ilvl = numPr.find(qn("w:ilvl")).get(qn("w:val"))
        numid = numPr.find(qn("w:numId")).get(qn("w:val"))
        return ilvl, numid

    l1, n1 = linked("Heading 1")
    l2, n2 = linked("Heading 2")
    l3, n3 = linked("Heading 3")
    assert (l1, l2, l3) == ("0", "1", "2")       # standard Word mapping
    assert n1 == n2 == n3                          # all the same list
    assert any("1 / 1.1 / 1.1.1" in a for a in result["applied"])


def test_headings_not_numbered_when_not_requested():
    from docx import Document
    from docx.oxml.ns import qn
    data = _thesis_with_headings()
    profile = {"document_structure": {"enabled": True, "headings": {
        "h1": {"size_pt": 15}}}}          # sizes only, numbered omitted (falsy)
    result = docformat.apply_profile_to_docx(data, profile)
    doc = Document(io.BytesIO(result["data"]))
    pPr = doc.styles["Heading 1"].element.find(qn("w:pPr"))
    numPr = pPr.find(qn("w:numPr")) if pPr is not None else None
    assert numPr is None                  # no numbering added
    assert docformat.text_fingerprint(result["data"]) == HEADING_PROSE


@pytest.fixture
def client(tmp_path):
    app = create_app(tmp_path / "fmt.db")
    c = TestClient(app); c.headers.update({"Authorization": f"Bearer {app.state.token}"})
    return c


def test_api_format_with_inline_profile(client):
    data = _thesis_docx()
    r = client.post("/v1/documents/format", json={
        "data_b64": base64.b64encode(data).decode(), "profile": PROFILE})
    assert r.status_code == 200
    body = r.json()
    assert body["applied"]
    # returned file is a real docx with identical prose
    out = base64.b64decode(body["data_b64"])
    assert docformat.text_fingerprint(out) == PROSE


def test_api_format_with_saved_profile_id(client):
    pid = client.post("/v1/style-profiles/build", json={
        "name": "Test Uni", "based_on_csl": "apa", "font_family": "Cambria",
        "font_size_pt": 12, "margin_cm": 3.0}).json()["id"]
    data = _thesis_docx()
    r = client.post("/v1/documents/format", json={
        "data_b64": base64.b64encode(data).decode(), "profile_id": pid})
    assert r.status_code == 200 and r.json()["applied"]


def test_api_format_bad_base64_is_400(client):
    r = client.post("/v1/documents/format", json={"data_b64": "!!!not base64!!!", "profile": PROFILE})
    assert r.status_code == 400


def test_api_format_missing_profile_is_400(client):
    data = base64.b64encode(_thesis_docx()).decode()
    r = client.post("/v1/documents/format", json={"data_b64": data})
    assert r.status_code == 400


def test_binding_margin_and_caption_style():
    from docx import Document
    from docx.shared import Cm, Pt
    import io as _io
    doc = Document()
    doc.add_heading("H", level=1); doc.add_paragraph("body")
    doc.add_paragraph("Figure 1.", style="Caption")
    buf = _io.BytesIO(); doc.save(buf)
    prof = {"paragraph_style": {"font_family": "Georgia", "font_size_pt": 12},
            "document_structure": {"enabled": True, "margin_cm": 2.54,
                "binding_margin_cm": 3.81, "binding_side": "left",
                "caption": {"size_pt": 10, "italic": True}}}
    r = docformat.apply_profile_to_docx(buf.getvalue(), prof)
    d = Document(io.BytesIO(r["data"]))
    s = d.sections[0]
    assert abs(s.left_margin.cm - 3.81) < 0.02 and abs(s.right_margin.cm - 2.54) < 0.02
    cs = d.styles["Caption"]
    assert cs.font.size == Pt(10) and cs.font.italic is True
    assert docformat.text_fingerprint(r["data"]) == docformat.text_fingerprint(buf.getvalue())


def test_headings_numbered_creates_numbering_part_when_absent():
    from docx import Document
    import io as _io
    doc = Document()            # fresh doc: no numbering part at all
    doc.add_heading("Intro", level=1)
    buf = _io.BytesIO(); doc.save(buf)
    prof = {"document_structure": {"enabled": True, "headings": {"numbered": True}}}
    r = docformat.apply_profile_to_docx(buf.getvalue(), prof)
    assert any("Heading numbering →" in a for a in r["applied"])   # applied, not skipped
    d = Document(io.BytesIO(r["data"]))
    from docx.oxml.ns import qn
    ppr = d.styles["Heading 1"].element.find(qn("w:pPr"))
    assert ppr is not None and ppr.find(qn("w:numPr")) is not None


def test_binding_margin_is_orientation_aware():
    from docx import Document
    from docx.enum.section import WD_ORIENT, WD_SECTION
    import io as _io
    d = Document()
    d.add_heading("Ch1", level=1)
    sec = d.add_section(WD_SECTION.NEW_PAGE); sec.orientation = WD_ORIENT.LANDSCAPE
    w, h = sec.page_height, sec.page_width; sec.page_width, sec.page_height = w, h
    d.add_heading("Wide", level=1)
    buf = _io.BytesIO(); d.save(buf)
    prof = {"document_structure": {"enabled": True, "margin_cm": 2.54,
            "binding_margin_cm": 3.81, "binding_side": "left", "binding_side_landscape": "bottom"}}
    r = docformat.apply_profile_to_docx(buf.getvalue(), prof)
    d2 = Document(io.BytesIO(r["data"]))
    port, land = d2.sections[0], d2.sections[1]
    assert abs(port.left_margin.cm - 3.81) < 0.02 and abs(port.bottom_margin.cm - 2.54) < 0.02
    assert abs(land.bottom_margin.cm - 3.81) < 0.02 and abs(land.left_margin.cm - 2.54) < 0.02
