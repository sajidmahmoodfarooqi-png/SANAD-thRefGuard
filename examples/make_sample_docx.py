#!/usr/bin/env python3
"""Generate examples/sample-thesis.docx — a short practice manuscript for trying SANAD.

It is deliberately plain prose (no pre-baked SANAD fields): you open it in Word with the
SANAD add-in, insert citations from ``sample-library.ris`` where marked, and watch the
Integrity check catch the planted problems. Run:  python examples/make_sample_docx.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "sample-thesis.docx"


def main():
    d = Document()
    d.styles["Normal"].font.name = "Calibri"
    d.styles["Normal"].font.size = Pt(11)

    t = d.add_heading("Planned Human Settlements and Urban Ecological Resilience", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = d.add_paragraph("A sample manuscript for trying SANAD — the RefGuard")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    note = d.add_paragraph()
    r = note.add_run("How to use this file:  ")
    r.bold = True
    note.add_run(
        "First import examples/sample-library.ris into SANAD. Then, in Word with the SANAD "
        "add-in open, place your cursor at each 【insert citation …】 marker below, switch to the "
        "add-in's Insert tab, search for the named source, and insert it. When you've inserted "
        "them all, open the Integrity tab and press Check this document. The section at the end "
        "tells you what each planted problem should trigger.")

    d.add_heading("1. Introduction", level=1)
    d.add_paragraph(
        "Rapid, unplanned urban expansion places mounting pressure on the ecological systems that "
        "cities depend upon. Resilient water networks are increasingly treated as a prerequisite "
        "for sustainable settlement growth 【insert citation: Alvarez & Okonkwo, resilient urban "
        "water networks】. A broader planning tradition frames these questions within ecological "
        "limits 【insert citation: Harding, Foundations of Ecological Planning】.")

    d.add_heading("2. Equity and the urban environment", level=1)
    d.add_paragraph(
        "The distribution of green infrastructure is rarely even, and spatial equity has become a "
        "central concern for planners 【insert citation: Nakamura, Spatial equity in urban green "
        "infrastructure】. Where growth is informal, remote sensing offers a practical way to track "
        "change over time 【insert citation: Costa et al., Remote sensing of informal settlement "
        "growth】.")

    d.add_heading("3. Methods", level=1)
    d.add_paragraph(
        "Participatory mapping complements satellite analysis by grounding it in local knowledge "
        "【insert citation: Bianchi, Participatory mapping methods】. For the statistical treatment "
        "of field observations we follow standard practice 【insert citation: Osei, Statistics for "
        "Field Researchers】.")

    d.add_heading("4. A note for testing the integrity checks", level=1)
    steps = [
        ("R1 — Year mismatch", "Insert the Alvarez & Okonkwo citation, then edit the visible "
         "in-text year to 2018. SANAD should flag it against the library's 2019."),
        ("R6 — Possible duplicate", "sample-library.ris contains the Alvarez & Okonkwo article "
         "twice (one with a DOI, one without). Open Library health in the app — it should offer "
         "them as a likely duplicate to review (never merged automatically)."),
        ("R7 — Ambiguous author/year", "The library has two 2020 articles by Nakamura with no a/b "
         "suffix. Cite both and SANAD should ask you to disambiguate."),
        ("R2 — Venue looks wrong", "'Harding, Foundations of Ecological Planning' is a book; the "
         "library also holds a review of it in Planning Review. If a book citation resolves to the "
         "review venue, SANAD flags the mismatch."),
        ("R8 — Cited out of context", "In the Methods section, deliberately insert the Delacroix "
         "'Land-use change and biodiversity loss' citation next to a sentence about statistics. "
         "With the semantic check on, SANAD should note the weak match and suggest a closer source."),
    ]
    for title, body in steps:
        p = d.add_paragraph(style="List Bullet")
        p.add_run(title + ".  ").bold = True
        p.add_run(body)

    foot = d.add_paragraph()
    fr = foot.add_run("SANAD only ever writes inside its own citation and reference-list fields — "
                      "it never edits your prose.")
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    d.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
